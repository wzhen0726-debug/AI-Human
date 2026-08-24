"""结论文档骨架: 数字人模型处理阶段性结论(给部门经理).
结构: 输入→AI处理流程→输出对比→关键结论. 图片素材渲染后填充."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SZRY = r"E:\WangZhen_Project\AI\ShuZiRen\Hermes\SZRYanJiu"
OUT_DIR = os.path.join(SZRY, "v3_QuadRemesher_交付", "汇报素材")
os.makedirs(OUT_DIR, exist_ok=True)
OUT_DOC = os.path.join(OUT_DIR, "数字人模型处理_阶段性结论.docx")

doc = Document()
# 默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '微软雅黑'
        r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h

def P(text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r.font.size = Pt(size)
    r.bold = bold
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r.font.size = Pt(10.5)
    return p

def img(path, width=Inches(5.5), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in c.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ============ 标题 ============
t = doc.add_heading('数字人模型处理工具 — 阶段性结论', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('输入 → AI处理 → 输出对比（2026年8月）')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ============ 一、一句话结论 ============
H('一、核心结论', 1)
P('一张 AI 生成的照片级高模（约 190万 面、全三角形、无拓扑规范、无法直接用于动画），'
  '经本工具链 6 步自动处理，产出可绑定、可动画、带 4K 贴图的低模（头部 14.3万面全四边形），'
  '全流程无人工建模操作，各步骤均为脚本自动化。', bold=True)
P('')
P('当前已完成并验证：高模修复（01）→ 眼窝与眼球（01A）→ 自动重拓扑（02）→ 自动UV（03）→ 4K纹理烘焙（04）。'
  '剩余骨骼绑定（05）与GLB导出（06）已有方案，待接入。')

# ============ 二、输入 ============
H('二、输入：一张照片生成的 3D 高模', 1)
P('输入来源：单张人像照片 → Tripo AI 生成 3D 模型（GLB 格式）。')
bullet('面数：约 190万 面（全三角形，本例实测 188.6万面 / 96.9万顶点），远超实时渲染承受能力')
bullet('拓扑：全三角形、无规律，眼窝/鼻孔等孔洞破损，无法做表情动画')
bullet('结构：无眼窝凹陷、无独立眼球、无法线信息，直接使用会穿帮')
img(os.path.join(OUT_DIR, '输入_原始Tripo高模_front.png'), caption='输入：Tripo AI 生成的原始高模（正面）')
img(os.path.join(OUT_DIR, '输入_原始Tripo高模_side.png'), caption='输入：原始高模（侧面）')

# ============ 三、AI 处理流程 ============
H('三、AI 处理流程（6 步自动化）', 1)
P('每一步都有独立脚本、日志与验收截图，可单步重跑、可追溯。')

steps = [
    ('01 高模修复', '旋转归位、焊接重复顶点、非流形修复、黏连检测；'
     '局部异常（腹部孔洞=法线朝内、胸口凸起）按解剖判断修复'),
    ('01A 眼窝与眼球', '眼窝：程序化挖出碗形眼窝（用户GUI验收，无凸脊、接缝平滑）；'
     '眼球：按解剖规律自动摆放（角膜贴开口平面+虹膜底贴下睑），可中文面板换色/微调'),
    ('02 自动重拓扑', 'QuadRemesher 自动生成全四边形低模（14.3万面、100%四边形）；'
     '验证：眼窝开口边缘偏差<0.7mm，可见结构完整保留'),
    ('03 自动UV', '自动展开UV，接缝少、无碎岛、无扭曲'),
    ('04 纹理烘焙', '高模细节烘焙为 4K 颜色图+法线图，低模获得高模观感'),
    ('05/06 绑定与导出', 'Mixamo 骨骼绑定 + GLB 导出（方案已定，待接入）'),
]
for name, desc in steps:
    p = doc.add_paragraph()
    r = p.add_run(name + '　')
    r.bold = True; r.font.name = '微软雅黑'
    r2 = p.add_run(desc)
    r2.font.name = '微软雅黑'; r2.font.size = Pt(10.5)

# ============ 四、输出对比 ============
H('四、输出对比', 1)
H('4.1 模型质量对比', 2)
tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Light Grid Accent 1'
rows = [
    ('指标', '输入（原始高模）', '输出（低模）'),
    ('面数', '188.6万面（全三角）', '14.3万面（全四边形）'),
    ('眼窝', '无（平面破损）', '程序化碗形，开口边缘保留<0.7mm偏差'),
    ('眼球', '无', '自动解剖定位，可换19色、可微调'),
    ('贴图', '无规范贴图', '4K 颜色+法线烘焙'),
]
for i, row in enumerate(rows):
    for j, cell in enumerate(row):
        c = tbl.cell(i, j)
        c.text = cell
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = '微软雅黑'
                r.font.size = Pt(9.5)
                if i == 0: r.bold = True

H('4.2 视觉对比', 2)
img(os.path.join(OUT_DIR, '修复后高模(含眼窝)_front.png'), caption='处理中高模：眼窝已成型（正面）')
img(os.path.join(OUT_DIR, '低模_带眼球_front.png'), caption='输出低模+眼球（正面）')
img(os.path.join(OUT_DIR, '低模_wire_front.png'), caption='输出低模线框：全四边形拓扑（正面）')

# ============ 五、关键技术结论 ============
H('五、关键技术结论', 1)
bullet('眼窝工作价值已验证：重拓扑后眼窝开口边缘偏差<0.7mm，可见结构完整保留；'
       '被磨平的碗内深处本就被眼球遮挡，不影响外观')
bullet('眼球定位规律可复用：角膜贴开口平面+虹膜底贴下睑，换眼球模型自动适配，'
       '用户GUI微调可保存回管线')
bullet('高模修复是地基：焊接/非流形/黏连修复后，重拓扑才能 52秒完成 100%四边形')
bullet('全流程脚本化：每步可单步重跑、有日志与验收截图，错误可追溯')

doc.save(OUT_DOC)
print(f"骨架已生成: {OUT_DOC}")
print("待补: 渲染图填充")
